# -*- coding: utf-8 -*-
# Consolidated kundali generator with D1/D9 rules and markers (library-only; no demo code)

from docx import Document
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _normalize_chalit_arrays(begins_sid, mids_sid):
    """Normalize begins_sid/mids_sid to 1-based dicts for safe access"""
    def _norm(x):
        if isinstance(x, dict):
            return x
        elif isinstance(x, (list, tuple)) and len(x) >= 12:
            return {i+1: x[i] for i in range(12)}
        else:
            return {i: 0 for i in range(1, 13)}  # fallback
    
    return _norm(begins_sid), _norm(mids_sid)

def detect_conjunctions_in_house(planets, max_degrees=6.0):
    """Detect conjunctions between planets in the same house."""
    conjunctions = []
    for i in range(len(planets)):
        for j in range(i + 1, len(planets)):
            p1, p2 = planets[i], planets[j]
            if (isinstance(p1, dict) and 'lon' in p1 and 
                isinstance(p2, dict) and 'lon' in p2):
                lon1, lon2 = p1['lon'] % 360, p2['lon'] % 360
                sep = min(abs(lon1 - lon2), 360 - abs(lon1 - lon2))
                if sep <= max_degrees:
                    conjunctions.append({
                        'planet1_idx': i, 'planet2_idx': j,
                        'separation_degrees': int(round(sep))
                    })
    return conjunctions

def draw_conjunction_arrow(cx1, cy1, cx2, cy2, degrees, color="#FF6600"):
    """Generate VML elements for conjunction arrow with degree label."""
    mid_x, mid_y = (cx1 + cx2) / 2, (cy1 + cy2) / 2
    elements = []
    
    # Double-arrowed line
    elements.append(f'''<v:line style="position:absolute;left:0pt;top:0pt;width:1pt;height:1pt;z-index:6" 
                        from="{cx1}pt,{cy1}pt" to="{cx2}pt,{cy2}pt" 
                        strokecolor="{color}" strokeweight="2.0pt">
                        <v:stroke endarrow="classic" startarrow="classic"/>
                    </v:line>''')
    
    # Degree label  
    elements.append(f'''<v:rect style="position:absolute;left:{mid_x-8}pt;top:{mid_y-6}pt;width:16pt;height:12pt;z-index:7" 
                        strokecolor="{color}" strokeweight="0.5pt" fillcolor="#FFF8DC">
                        <v:textbox inset="1pt,1pt,1pt,1pt">
                            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:p><w:pPr><w:jc w:val="center"/></w:pPr>
                                <w:r><w:rPr><w:sz w:val="14"/><w:b/></w:rPr>
                                <w:t>{degrees}°</w:t></w:r></w:p>
                            </w:txbxContent>
                        </v:textbox>
                    </v:rect>''')
    return elements

HN_ABBR = {
    'Su': 'सू', 'Mo': 'चं', 'Ma': 'मं', 'Me': 'बु',
    'Ju': 'गु', 'Ve': 'शु', 'Sa': 'श', 'Ra': 'रा', 'Ke': 'के'
}

SELF_SIGNS = {
    'Su': {5}, 'Mo': {4}, 'Ma': {1, 8}, 'Me': {3, 6}, 'Ju': {9, 12},
    'Ve': {2, 7}, 'Sa': {10, 11}, 'Ra': set(), 'Ke': set()
}

EXALT_SIGN = {'Su': 1, 'Mo': 2, 'Ma': 10, 'Me': 6, 'Ju': 4, 'Ve': 12, 'Sa': 7}
DEBIL_SIGN = {p: ((s + 5) % 12) + 1 for p, s in EXALT_SIGN.items()}
COMB_ORB = {'Mo': 12, 'Ma': 17, 'Me': 12, 'Ju': 11, 'Ve': 10, 'Sa': 15, 'Ra': 0, 'Ke': 0}

UP_ARROW, DOWN_ARROW, COMBUST, VARG_SQ = '↑', '↓', '^', '◱'

def _sep_deg(a, b):
    d = abs((a - b) % 360.0)
    return d if d <= 180 else 360 - d

def _rasi_sign(lon_sid):
    return int(lon_sid // 30) + 1

def navamsa_sign_from_lon_sid(lon_sid):
    rasi = int(lon_sid // 30) + 1
    part = int((lon_sid % 30) // (30/9.0))
    return ((rasi - 1) * 9 + part) % 12 + 1

def _is_combust_d1(code, sidelons):
    if code not in COMB_ORB or COMB_ORB[code] == 0: return False
    return _sep_deg(sidelons[code], sidelons['Su']) <= COMB_ORB[code]

def _is_combust_d9_same_nsign(code, sidelons):
    if code in ('Su','Ra','Ke'): return False
    sun_n = navamsa_sign_from_lon_sid(sidelons['Su'])
    pl_n  = navamsa_sign_from_lon_sid(sidelons[code])
    return sun_n == pl_n

def build_rasi_house_planets(sidelons, lagna_sign, begins_sid=None):
    """Build rasi house planets with longitude data for precise positioning."""
    house_map = {i: [] for i in range(1, 13)}
    for code in HN_ABBR.keys():
        lon = sidelons[code]; rasi = _rasi_sign(lon)
        
        # Use Chalit house system if available, otherwise fall back to rasi houses
        if begins_sid:
            from chalit_kundali_renderer import house_index_for_lon_chalit
            h = house_index_for_lon_chalit(lon, begins_sid)
        else:
            h = ((rasi - lagna_sign) % 12) + 1
            
        is_self = rasi in SELF_SIGNS.get(code, set())
        is_ex   = (code not in ('Ra','Ke')) and (EXALT_SIGN.get(code) == rasi)
        is_de   = (code not in ('Ra','Ke')) and (DEBIL_SIGN.get(code) == rasi)
        is_cb   = _is_combust_d1(code, sidelons)
        is_vg   = (rasi == navamsa_sign_from_lon_sid(lon))
        base = HN_ABBR[code]; disp = base
        if is_ex: disp += UP_ARROW
        if is_de: disp += DOWN_ARROW
        if is_cb: disp += COMBUST
        house_map[h].append({
            'txt': base,
            'disp': disp,
            'lon': lon,  # Include longitude for precise positioning
            'flags': {'self':is_self,'exalt':is_ex,'debil':is_de,'comb':is_cb,'varg':is_vg}
        })
    return house_map

def build_navamsa_house_planets(sidelons, nav_lagna_sign):
    house_map = {i: [] for i in range(1, 13)}
    for code in HN_ABBR.keys():
        lon = sidelons[code]; nsign = navamsa_sign_from_lon_sid(lon)
        h = ((nsign - nav_lagna_sign) % 12) + 1
        is_self = nsign in SELF_SIGNS.get(code, set())
        is_ex   = (code not in ('Ra','Ke')) and (EXALT_SIGN.get(code) == nsign)
        is_de   = (code not in ('Ra','Ke')) and (DEBIL_SIGN.get(code) == nsign)
        is_cb   = _is_combust_d9_same_nsign(code, sidelons)
        is_vg   = (_rasi_sign(lon) == nsign)
        base = HN_ABBR[code]; disp = base
        if is_ex: disp += UP_ARROW
        if is_de: disp += DOWN_ARROW
        if is_cb: disp += COMBUST
        house_map[h].append({'txt': base,'disp':disp,
            'flags':{'self':is_self,'exalt':is_ex,'debil':is_de,'comb':is_cb,'varg':is_vg}})
    return house_map

def calculate_precise_planet_coordinates(house_planets, begins_sid, mids_sid, size_pt=220):
    """
    Calculate precise coordinates for planets within houses based on Chalit boundaries.
    
    Args:
        house_planets: Dict mapping house -> list of planet dictionaries with 'lon' (longitude)
        begins_sid: List of house begin longitudes from Chalit table (1-indexed)  
        mids_sid: List of house mid longitudes from Chalit table (1-indexed)
        size_pt: Size of the kundali in points
        
    Returns:
        Dict mapping house -> list of (x, y) coordinates for planets in that house
    """
    S = size_pt
    
    # Define house center coordinates and boundaries for North Indian kundali
    house_centers = {
        1: (S/2, S/8),     # Top center
        2: (3*S/4, S/4),   # Top right  
        3: (7*S/8, S/2),   # Right center
        4: (3*S/4, 3*S/4), # Bottom right
        5: (S/2, 7*S/8),   # Bottom center
        6: (S/4, 3*S/4),   # Bottom left
        7: (S/8, S/2),     # Left center
        8: (S/4, S/4),     # Top left
        9: (S/2, S/2.7),   # Inner top center
        10: (S/1.35, S/2), # Inner right center
        11: (S/2, S/1.35), # Inner bottom center  
        12: (S/2.7, S/2)   # Inner left center
    }
    
    # Define house boundaries for positioning (start, mid, end positioning areas)
    house_bounds = {
        1: {'start': (S/2, S/16), 'mid': (S/2, S/8), 'end': (S/2, 3*S/16)},
        2: {'start': (5*S/8, S/4), 'mid': (3*S/4, S/4), 'end': (7*S/8, S/4)},
        3: {'start': (7*S/8, 3*S/8), 'mid': (7*S/8, S/2), 'end': (7*S/8, 5*S/8)},
        4: {'start': (7*S/8, 3*S/4), 'mid': (3*S/4, 3*S/4), 'end': (5*S/8, 3*S/4)},
        5: {'start': (S/2, 15*S/16), 'mid': (S/2, 7*S/8), 'end': (S/2, 13*S/16)},
        6: {'start': (3*S/8, 3*S/4), 'mid': (S/4, 3*S/4), 'end': (S/8, 3*S/4)},
        7: {'start': (S/8, 5*S/8), 'mid': (S/8, S/2), 'end': (S/8, 3*S/8)},
        8: {'start': (S/8, S/4), 'mid': (S/4, S/4), 'end': (3*S/8, S/4)},
        9: {'start': (S/2, S/3.5), 'mid': (S/2, S/2.7), 'end': (S/2, S/2.2)},
        10: {'start': (S/1.6, S/2), 'mid': (S/1.35, S/2), 'end': (S/1.15, S/2)},
        11: {'start': (S/2, S/2.2), 'mid': (S/2, S/1.35), 'end': (S/2, S/1.1)},
        12: {'start': (S/3.5, S/2), 'mid': (S/2.7, S/2), 'end': (S/2.2, S/2)}
    }
    
    planet_coords = {}
    
    for house_num in range(1, 13):
        if house_num not in house_planets or not house_planets[house_num]:
            continue
            
        house_planets_list = house_planets[house_num]
        planet_positions = []
        
        # Normalize Chalit arrays to 1-based dicts for safe access
        B, M = _normalize_chalit_arrays(begins_sid, mids_sid)
        
        # Get Chalit boundaries for this house (now safely normalized)
        begin_lon = B[house_num] % 360
        mid_lon = M[house_num] % 360
        end_lon = B[(house_num % 12) + 1] % 360
        
        bounds = house_bounds.get(house_num, house_centers[house_num])
        
        for planet_info in house_planets_list:
            if not isinstance(planet_info, dict) or 'lon' not in planet_info:
                # Default positioning for planets without longitude data
                planet_positions.append(house_centers[house_num])
                continue
                
            planet_lon = planet_info['lon']
            
            # Normalize longitudes to 0-360 range
            begin_lon = begin_lon % 360
            mid_lon = mid_lon % 360  
            end_lon = end_lon % 360
            planet_lon = planet_lon % 360
            
            # Determine if planet is in start-to-mid or mid-to-end segment
            start_x, start_y = bounds['start']
            mid_x, mid_y = bounds['mid']  
            end_x, end_y = bounds['end']
            
            # Calculate position ratio within the house
            if _angle_between(begin_lon, planet_lon, mid_lon):
                # Planet is between start and mid bhava
                ratio = _angle_ratio(begin_lon, planet_lon, mid_lon)
                x = start_x + ratio * (mid_x - start_x)
                y = start_y + ratio * (mid_y - start_y)
            else:
                # Planet is between mid and end bhava  
                ratio = _angle_ratio(mid_lon, planet_lon, end_lon)
                x = mid_x + ratio * (end_x - mid_x)
                y = mid_y + ratio * (end_y - mid_y)
                
            planet_positions.append((x, y))
            
        planet_coords[house_num] = planet_positions
        
    return planet_coords

def _angle_between(start, angle, end):
    """Check if angle is between start and end, handling 0/360 wrap-around."""
    start = start % 360
    end = end % 360
    angle = angle % 360
    
    if start <= end:
        return start <= angle <= end
    else:  # Crosses 0/360 boundary
        return angle >= start or angle <= end

def _angle_ratio(start, angle, end):
    """Calculate ratio (0.0 to 1.0) of where angle falls between start and end."""
    start = start % 360
    end = end % 360
    angle = angle % 360
    
    if start <= end:
        if end - start == 0:
            return 0.5  # Avoid division by zero
        return (angle - start) / (end - start)
    else:  # Crosses 0/360 boundary
        total_range = (360 - start) + end
        if total_range == 0:
            return 0.5
        if angle >= start:
            return (angle - start) / total_range
        else:
            return ((360 - start) + angle) / total_range

def kundali_single_box(size_pt=220, house_planets=None, begins_sid=None, mids_sid=None, color="#FF6600"):
    S=size_pt; w,h=36,28
    
    # Generate color variants for theming
    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def rgb_to_hex(rgb):
        return '#{:02x}{:02x}{:02x}'.format(int(rgb[0]), int(rgb[1]), int(rgb[2]))
    
    def lighten_color(hex_color, factor=0.7):
        r, g, b = hex_to_rgb(hex_color)
        r = int(r + (255 - r) * factor)
        g = int(g + (255 - g) * factor)
        b = int(b + (255 - b) * factor)
        return rgb_to_hex((r, g, b))
    
    def darken_color(hex_color, factor=0.3):
        r, g, b = hex_to_rgb(hex_color)
        r = int(r * (1 - factor))
        g = int(g * (1 - factor))
        b = int(b * (1 - factor))
        return rgb_to_hex((r, g, b))
    
    # Create color variants for different elements
    stroke_color = color  # Use original color for borders
    fill_color = lighten_color(color, 0.7)   # Light variant for fill
    
    rect=f'<v:rect style="width:{S}pt;height:{S}pt" strokecolor="{stroke_color}" strokeweight="2.5pt" fillcolor="{fill_color}"/>'
    diag1=f'<v:line from="0,0" to="{S}pt,{S}pt" strokecolor="{stroke_color}" strokeweight="1.5pt"/>'
    diag2=f'<v:line from="{S}pt,0" to="0,{S}pt" strokecolor="{stroke_color}" strokeweight="1.5pt"/>'
    mid1=f'<v:line from="{S/2}pt,0" to="{S}pt,{S/2}pt" strokecolor="{stroke_color}" strokeweight="1.5pt"/>'
    mid2=f'<v:line from="{S}pt,{S/2}pt" to="{S/2}pt,{S}pt" strokecolor="{stroke_color}" strokeweight="1.5pt"/>'
    mid3=f'<v:line from="{S/2}pt,{S}pt" to="0,{S/2}pt" strokecolor="{stroke_color}" strokeweight="1.5pt"/>'
    mid4=f'<v:line from="0,{S/2}pt" to="{S/2}pt,0" strokecolor="{stroke_color}" strokeweight="1.5pt"/>'
    
    # Default fixed coordinates (fallback when precise positioning is not available)
    coords={1:(S/2,S/8),2:(3*S/4,S/4),3:(7*S/8,S/2),4:(3*S/4,3*S/4),
            5:(S/2,7*S/8),6:(S/4,3*S/4),7:(S/8,S/2),8:(S/4,S/4),
            9:(S/2,S/2.7),10:(S/1.35,S/2),11:(S/2,S/1.35),12:(S/2.7,S/2)}
    
    # Use precise coordinates if Chalit data is available
    if begins_sid and mids_sid and house_planets:
        try:
            precise_coords = calculate_precise_planet_coordinates(house_planets, begins_sid, mids_sid, size_pt)
        except Exception as e:
            print(f"DEBUG: Error in precise coordinate calculation: {e}")
            precise_coords = {}
    else:
        precise_coords = {}
    
    nums={i:str(i) for i in range(1,13)}
    glyph_w,gap=8.0,3.0;r_circle,sq_size=5.2,4.2;y_off=4.0
    boxes=[]; overlays=[]; conjunction_lines=[]
    
    for k in range(1, 13):
        # Use precise coordinates if available, otherwise fall back to fixed coords
        if k in precise_coords and precise_coords[k]:
            house_coord_list = precise_coords[k]
        else:
            # Default house center for planets without precise positioning
            default_coord = coords[k]
            house_coord_list = [default_coord] if k in house_planets and house_planets[k] else []
        
        items=house_planets.get(k,[]) if house_planets else []
        if items and not isinstance(items[0],dict): items=[{'txt':s,'disp':s,'flags':{}} for s in items]
        texts=[it.get('disp',it.get('txt','')) for it in items]; planets_text=" ".join(texts)
        
        # Use house center for textbox (house number)
        text_x, text_y = coords[k]
        content=f'<w:r><w:t>{nums[k]}</w:t></w:r>'
        # Only add combined planets_text if no precise coordinates available  
        if planets_text and not house_coord_list: 
            content+=f'<w:br/><w:r><w:t>{planets_text}</w:t></w:r>'
        left,top=text_x-w/2,text_y-h/2
        
        # Remove strokecolor border from textboxes as requested
        boxes.append(f'<v:rect style="position:absolute;left:{left}pt;top:{top}pt;width:{w}pt;height:{h}pt;z-index:5" strokecolor="none" fillcolor="none"><v:textbox inset="0,0,0,0"><w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:p><w:pPr><w:jc w:val="center"/></w:pPr>{content}</w:p></w:txbxContent></v:textbox></v:rect>')
        
        # Ensure house_coord_list length matches items length
        if items:
            # Detect conjunctions within this house (planets within 6 degrees)  
            conjunctions = detect_conjunctions_in_house(items, 6.0)
            
            # Ensure coordinate list matches items length
            if not house_coord_list:
                house_coord_list = [coords[k]] * len(items)  # Default to house center
            elif len(house_coord_list) < len(items):
                # Pad with last coordinate or house center
                last_coord = house_coord_list[-1] if house_coord_list else coords[k]
                while len(house_coord_list) < len(items):
                    house_coord_list.append(last_coord)
            
            for idx, (item, coord) in enumerate(zip(items, house_coord_list)):
                cx, cy = coord
                flags = item.get('flags', {})
                
                # Render planet glyph text at precise coordinates
                planet_text = item.get('disp', item.get('txt', ''))
                if planet_text:
                    boxes.append(f'''<v:rect style="position:absolute;left:{cx-glyph_w/2}pt;top:{cy-6}pt;width:{glyph_w}pt;height:12pt;z-index:8" 
                                     strokecolor="none" fillcolor="none">
                                     <v:textbox inset="0,0,0,0">
                                         <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                             <w:p><w:pPr><w:jc w:val="center"/></w:pPr>
                                             <w:r><w:rPr><w:sz w:val="16"/><w:b/></w:rPr>
                                             <w:t>{planet_text}</w:t></w:r></w:p>
                                         </w:txbxContent>
                                     </v:textbox>
                                 </v:rect>''')
                
                # Draw planetary status markers using correct flag keys
                if flags.get('self'):
                    overlays.append(f'<v:oval style="position:absolute;left:{cx-r_circle}pt;top:{cy-6.0}pt;width:{2*r_circle}pt;height:{2*r_circle}pt;z-index:4" strokecolor="black" strokeweight="1pt" fillcolor="none"/>')
                if flags.get('exalt'):  # Fixed: 'exalt' not 'exalted'
                    overlays.append(f'<v:oval style="position:absolute;left:{cx-r_circle}pt;top:{cy-6.0}pt;width:{2*r_circle}pt;height:{2*r_circle}pt;z-index:4" strokecolor="green" strokeweight="1.5pt" fillcolor="none"/>')
                if flags.get('debil'):  # Fixed: 'debil' not 'debilitated'
                    overlays.append(f'<v:oval style="position:absolute;left:{cx-r_circle}pt;top:{cy-6.0}pt;width:{2*r_circle}pt;height:{2*r_circle}pt;z-index:4" strokecolor="red" strokeweight="1.5pt" fillcolor="none"/>')
                if flags.get('comb'):   # Fixed: 'comb' not 'combust'
                    overlays.append(f'<v:oval style="position:absolute;left:{cx-2}pt;top:{cy-2}pt;width:4pt;height:4pt;z-index:4" strokecolor="orange" strokeweight="1pt" fillcolor="orange"/>')
                if flags.get('varg'):
                    overlays.append(f'<v:rect style="position:absolute;left:{cx+4.2}pt;top:{cy-7.4}pt;width:{sq_size}pt;height:{sq_size}pt;z-index:4" strokecolor="black" strokeweight="1pt" fillcolor="none"/>')
            
            # Draw conjunction arrows and degree labels  
            for c in conjunctions:
                i, j = c['planet1_idx'], c['planet2_idx']
                deg = c['separation_degrees']
                if i < len(house_coord_list) and j < len(house_coord_list):
                    x1, y1 = house_coord_list[i]
                    x2, y2 = house_coord_list[j]
                    overlays.extend(draw_conjunction_arrow(x1, y1, x2, y2, deg, color))
    
    xml=f'<w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:o="urn:schemas-microsoft-com:office:office">{rect}{diag1}{diag2}{mid1}{mid2}{mid3}{mid4}{"".join(boxes)}{"".join(overlays)}</w:pict>'
    return parse_xml(xml)

def detect_conjunctions_in_house(planets, max_degrees=6.0):
    """
    Detect conjunctions between planets in the same house.
    Returns list of conjunction info with planet indices and degree separation.
    """
    conjunctions = []
    
    for i in range(len(planets)):
        for j in range(i + 1, len(planets)):
            planet1 = planets[i]
            planet2 = planets[j]
            
            # Check if both planets have longitude data
            if not (isinstance(planet1, dict) and 'lon' in planet1 and 
                    isinstance(planet2, dict) and 'lon' in planet2):
                continue
                
            lon1 = planet1['lon'] % 360
            lon2 = planet2['lon'] % 360
            
            # Calculate angular separation 
            sep = min(abs(lon1 - lon2), 360 - abs(lon1 - lon2))
            
            if sep <= max_degrees:
                conjunctions.append({
                    'planet1_idx': i,
                    'planet2_idx': j,
                    'separation_degrees': int(round(sep))
                })
    
    return conjunctions

def draw_conjunction_arrow(coords_list, conjunction_info, planets):
    """
    Draw double-arrowed line between conjunct planets with degree label.
    """
    elements = []
    
    idx1 = conjunction_info['planet1_idx']
    idx2 = conjunction_info['planet2_idx']
    degrees = conjunction_info['separation_degrees']
    
    if idx1 >= len(coords_list) or idx2 >= len(coords_list):
        return elements
        
    x1, y1 = coords_list[idx1]
    x2, y2 = coords_list[idx2]
    
    # Calculate midpoint for degree label
    mid_x = (x1 + x2) / 2
    mid_y = (y1 + y2) / 2
    
    # Adjust label position for very close planets (2-3 degrees)
    if degrees <= 3:
        # Place label above the planets
        label_y = min(y1, y2) - 8
    else:
        # Place label at midpoint
        label_y = mid_y - 4
    
    # Draw double-arrowed line using proper VML arrow syntax with themed colors
    elements.append(f'''<v:line style="position:absolute;left:0pt;top:0pt;width:1pt;height:1pt;z-index:6" 
                        from="{x1}pt,{y1}pt" to="{x2}pt,{y2}pt" 
                        strokecolor="{color}" strokeweight="2.0pt">
                        <v:stroke endarrow="classic" startarrow="classic" 
                                  endarrowwidth="narrow" startarrowwidth="narrow"/>
                    </v:line>''')
    
    # Create color variants for label styling
    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def rgb_to_hex(rgb):
        return '#{:02x}{:02x}{:02x}'.format(int(rgb[0]), int(rgb[1]), int(rgb[2]))
    
    def lighten_color(hex_color, factor=0.8):
        r, g, b = hex_to_rgb(hex_color)
        r = int(r + (255 - r) * factor)
        g = int(g + (255 - g) * factor)
        b = int(b + (255 - b) * factor)
        return rgb_to_hex((r, g, b))
    
    def darken_color(hex_color, factor=0.2):
        r, g, b = hex_to_rgb(hex_color)
        r = int(r * (1 - factor))
        g = int(g * (1 - factor))
        b = int(b * (1 - factor))
        return rgb_to_hex((r, g, b))
    
    # Use color variants for themed styling
    fill_color = lighten_color(color, 0.8)  # Very light for background
    text_color = darken_color(color, 0.2)   # Slightly darker for text
    
    # Add degree label with themed colors
    elements.append(f'''<v:rect style="position:absolute;left:{mid_x-10}pt;top:{label_y}pt;width:20pt;height:14pt;z-index:7" 
                        strokecolor="{color}" strokeweight="0.75pt" fillcolor="{fill_color}">
                        <v:textbox inset="1pt,1pt,1pt,1pt">
                            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:p><w:pPr><w:jc w:val="center"/></w:pPr>
                                <w:r><w:rPr><w:sz w:val="16"/><w:b/><w:color w:val="{text_color.lstrip('#')}"/></w:rPr>
                                <w:t>{degrees}°</w:t></w:r></w:p>
                            </w:txbxContent>
                        </v:textbox>
                    </v:rect>''')
    
    return elements

def add_kundali_to_doc(doc,title,house_planets,size_pt=220,begins_sid=None,mids_sid=None,color="#FF6600"):
    p_title=doc.add_paragraph(); p_title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p_title.add_run(title); r.bold=True; r.underline=True
    p=doc.add_paragraph(); p._p.append(kundali_single_box(size_pt,house_planets,begins_sid,mids_sid,color)); doc.add_paragraph('')

def render_kundalis_into_doc(doc,sidelons,lagna_sign,nav_lagna_sign,size_pt=220,begins_sid=None,mids_sid=None,color="#FF6600"):
    """Render kundalis with precise planetary positioning based on Chalit boundaries."""
    rasi_map=build_rasi_house_planets(sidelons,lagna_sign,begins_sid)
    nav_map=build_navamsa_house_planets(sidelons,nav_lagna_sign)  # Navamsa keeps traditional positioning
    add_kundali_to_doc(doc,'लग्न कुंडली',rasi_map,size_pt,begins_sid,mids_sid,color)
    add_kundali_to_doc(doc,'नवांश कुंडली',nav_map,size_pt,color=color)  # Navamsa without precise positioning